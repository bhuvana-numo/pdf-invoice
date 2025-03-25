from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def replace_text(doc, replacements):
    """Replaces placeholders in a Word document."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        para.text = para.text.replace(f"{{{{{key}}}}}", str(value))

def insert_images(doc, logo_path, qr_path):
    """Replaces image placeholders with actual images."""
    for para in doc.paragraphs:
        if "[IMAGE PLACEHOLDER]" in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(logo_path, width=Inches(1))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  

        elif "{{invoiceQR}}" in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(qr_path, width=Inches(1))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  

def fill_invoice_template(input_docx, output_docx, logo_path, qr_path, data):
    """Fills invoice template and saves output."""
    doc = Document(input_docx)
    replace_text(doc, data)
    insert_images(doc, logo_path, qr_path)
    doc.save(output_docx)
    print(f"Filled DOCX saved: {output_docx}")

if __name__ == "__main__":
    invoice_data = {
        "invoice": "INV-20250321",
        "stationAddress": "Mumbai, India",
        "stationGSTN": "GST12345678",
        "stationPAN": "PAN12345",
        "userName": "John Doe",
        "userType": "Corporate",
        "vehicleDetails": "Tesla Model 3",
        "accountAddress": "123 Business St, Delhi",
        "accountGSTN": "GST87654321",
        "accountPAN": "PAN98765",
        "createdOn": "21 March 2025",
        "stationName": "Jio BP Station",
        "cpid": "CP-001",
        "connectorType": "Type-2"
    }

    TEMPLATE_DOCX = "template_invoice.docx"
    FILLED_DOCX = "filled_invoice.docx"
    LOGO_PATH = "logo.png"
    QR_PATH = "invoice_qr.png"

    fill_invoice_template(TEMPLATE_DOCX, FILLED_DOCX, LOGO_PATH, QR_PATH, invoice_data)









